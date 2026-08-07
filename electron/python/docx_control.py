#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Word (.docx) control via python-docx — no COM required."""

import json
import os
import re
import shutil
import sys
import traceback


def debug_print(msg):
    print(msg, file=sys.stderr)


try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    debug_print("警告: python-docx 未安装")


class DocxController:
    def __init__(self):
        self.doc = None
        self.path = None

    def create_new(self):
        if not HAS_DOCX:
            return {"success": False, "error": "python-docx 未安装"}
        self.doc = Document()
        self.path = None
        return {"success": True, "message": "已创建空白文档"}

    def open_file(self, file_path):
        if not HAS_DOCX:
            return {"success": False, "error": "python-docx 未安装"}
        if not file_path or not os.path.isfile(file_path):
            return {"success": False, "error": f"文件不存在: {file_path}"}
        self.doc = Document(file_path)
        self.path = os.path.abspath(file_path)
        return {"success": True, "path": self.path, "paragraphs": len(self.doc.paragraphs)}

    def save_file(self, file_path=None):
        if not self.doc:
            return {"success": False, "error": "没有打开的文档"}
        target = file_path or self.path
        if not target:
            return {"success": False, "error": "未指定保存路径"}
        os.makedirs(os.path.dirname(os.path.abspath(target)) or ".", exist_ok=True)
        self.doc.save(target)
        self.path = os.path.abspath(target)
        return {"success": True, "path": self.path}

    def add_heading(self, text, level=1):
        if not self.doc:
            return {"success": False, "error": "没有打开的文档"}
        level = max(0, min(9, int(level or 1)))
        self.doc.add_heading(str(text or ""), level=level)
        return {"success": True}

    def add_paragraph(self, text, style=None):
        if not self.doc:
            return {"success": False, "error": "没有打开的文档"}
        p = self.doc.add_paragraph(str(text or ""))
        if style:
            try:
                p.style = style
            except Exception:
                pass
        return {"success": True}

    def set_paragraph(self, index, text):
        if not self.doc:
            return {"success": False, "error": "没有打开的文档"}
        idx = int(index)
        if idx < 0 or idx >= len(self.doc.paragraphs):
            return {"success": False, "error": f"段落索引越界: {idx}"}
        p = self.doc.paragraphs[idx]
        # keep first run style if possible
        if p.runs:
            p.runs[0].text = str(text or "")
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = str(text or "")
        return {"success": True, "index": idx}

    def replace_text(self, old, new, count=0):
        if not self.doc:
            return {"success": False, "error": "没有打开的文档"}
        old = str(old or "")
        new = str(new if new is not None else "")
        if not old:
            return {"success": False, "error": "old 不能为空"}
        replaced = 0
        limit = int(count or 0)

        def do_replace(text):
            nonlocal replaced
            if old not in text:
                return text
            if limit <= 0:
                n = text.count(old)
                replaced += n
                return text.replace(old, new)
            out = text
            while old in out and (limit <= 0 or replaced < limit):
                out = out.replace(old, new, 1)
                replaced += 1
            return out

        for p in self.doc.paragraphs:
            if old in p.text:
                full = do_replace(p.text)
                if p.runs:
                    p.runs[0].text = full
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = full
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old in cell.text:
                        for p in cell.paragraphs:
                            if old in p.text:
                                full = do_replace(p.text)
                                if p.runs:
                                    p.runs[0].text = full
                                    for r in p.runs[1:]:
                                        r.text = ""
                                else:
                                    p.text = full
        return {"success": True, "replaced": replaced}

    def get_outline(self, max_paragraphs=80):
        if not self.doc:
            return {"success": False, "error": "没有打开的文档"}
        items = []
        max_n = max(1, min(500, int(max_paragraphs or 80)))
        for i, p in enumerate(self.doc.paragraphs[:max_n]):
            style = p.style.name if p.style else ""
            text = (p.text or "").strip()
            if not text:
                continue
            level = None
            if style.lower().startswith("heading"):
                m = re.search(r"(\d+)", style)
                level = int(m.group(1)) if m else 1
            items.append({"index": i, "style": style, "level": level, "text": text[:200]})
        return {
            "success": True,
            "path": self.path,
            "paragraph_count": len(self.doc.paragraphs),
            "outline": items,
        }

    def fill_placeholders(self, mapping):
        """Replace {{key}} placeholders."""
        if not self.doc:
            return {"success": False, "error": "没有打开的文档"}
        if not isinstance(mapping, dict):
            return {"success": False, "error": "mapping 必须是对象"}
        total = 0
        for key, value in mapping.items():
            token = "{{" + str(key) + "}}"
            res = self.replace_text(token, str(value if value is not None else ""), 0)
            if res.get("success"):
                total += int(res.get("replaced") or 0)
        return {"success": True, "replaced": total}

    def close(self):
        self.doc = None
        self.path = None
        return {"success": True}


_controller = DocxController()


def run_script(script, goal=""):
    """Execute a python script with helpers in scope."""
    if not HAS_DOCX:
        return {"success": False, "error": "python-docx 未安装"}
    local_vars = {
        "Document": Document,
        "Pt": Pt,
        "RGBColor": RGBColor,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "controller": _controller,
        "doc": _controller.doc,
        "goal": goal or "",
        "os": os,
        "shutil": shutil,
        "json": json,
    }

    def save(path=None):
        return _controller.save_file(path)

    def open_doc(path):
        r = _controller.open_file(path)
        local_vars["doc"] = _controller.doc
        return r

    def create():
        r = _controller.create_new()
        local_vars["doc"] = _controller.doc
        return r

    local_vars.update({"save": save, "open_doc": open_doc, "create": create})
    try:
        exec(str(script or ""), {"__builtins__": __builtins__}, local_vars)
        if _controller.doc is not None:
            local_vars["doc"] = _controller.doc
        return {
            "success": True,
            "message": "脚本执行完成",
            "path": _controller.path,
            "goal": goal or "",
        }
    except Exception as e:
        return {"success": False, "error": str(e), "trace": traceback.format_exc()}


def copy_template(source, dest):
    if not source or not os.path.isfile(source):
        return {"success": False, "error": f"模板不存在: {source}"}
    dest = os.path.abspath(dest)
    os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)
    shutil.copy2(source, dest)
    return {"success": True, "path": dest}


def execute_command(action, args):
    args = args or {}
    try:
        if action in ("create_new", "create"):
            return _controller.create_new()
        if action in ("open_file", "open", "load_file", "load"):
            return _controller.open_file(args.get("file_path") or args.get("path"))
        if action in ("save_file", "save"):
            return _controller.save_file(args.get("file_path") or args.get("path"))
        if action == "add_heading":
            return _controller.add_heading(args.get("text"), args.get("level", 1))
        if action == "add_paragraph":
            return _controller.add_paragraph(args.get("text"), args.get("style"))
        if action == "set_paragraph":
            return _controller.set_paragraph(args.get("index", 0), args.get("text"))
        if action == "replace_text":
            return _controller.replace_text(args.get("old"), args.get("new"), args.get("count", 0))
        if action == "get_outline":
            return _controller.get_outline(args.get("max_paragraphs", 80))
        if action == "fill_placeholders":
            return _controller.fill_placeholders(args.get("mapping") or args.get("values") or {})
        if action == "close":
            return _controller.close()
        if action == "copy_template":
            return copy_template(args.get("source"), args.get("dest") or args.get("file_path"))
        if action == "run_script":
            return run_script(args.get("script"), args.get("goal", ""))
        return {"success": False, "error": f"未知 action: {action}"}
    except Exception as e:
        return {"success": False, "error": str(e), "trace": traceback.format_exc()}


def main():
    try:
        raw = sys.argv[1] if len(sys.argv) > 1 else sys.stdin.read()
        payload = json.loads(raw)
        action = payload.get("action")
        args = payload.get("args") or {}
        result = execute_command(action, args)
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
